from __future__ import annotations
import os
import time
import re
from dotenv import load_dotenv
import chainlit as cl

from docx import Document
from azure.ai.projects import AIProjectClient
from azure.identity import DefaultAzureCredential

# Load environment variables
load_dotenv()

# Azure Foundry setup
AIPROJECT_CONNECTION_STRING = os.getenv("AIPROJECT_CONNECTION_STRING")
project_client = AIProjectClient.from_connection_string(
    conn_str=AIPROJECT_CONNECTION_STRING, credential=DefaultAzureCredential()
)

AGENT_NAMES = ["AIS Requirements", "AIS Architect", "AIS Developer"]
AGENT_LOOKUP = {}


def get_agents_by_name(client, names: list[str]):
    found_agents = []
    try:
        agents = client.agents.list_agents()
        for name in names:
            match = next((a for a in agents.data if a.name == name), None)
            if match:
                found_agents.append(match)
            else:
                print(f"Agent named '{name}' not found.")
    except Exception as e:
        print(f"Error listing agents: {e}")
    return found_agents


def detect_target_agent(message: str) -> str | None:
    message = message.lower()
    if "requirement" in message:
        return "AIS Requirements"
    elif "architecture" in message:
        return "AIS Architect"
    elif "project" in message or "integration" in message:
        return "AIS Developer"
    return None


def should_generate_flow_doc(message: str) -> bool:
    trigger_phrases = [
        "generate integration flow",
        "create integration flow document",
        "integration flow doc",
        "integration flow documentation"
    ]
    return any(phrase in message.lower() for phrase in trigger_phrases)


def format_integration_flow_doc(title: str, content: str) -> Document:
    doc = Document()
    doc.add_heading(title or "Integration Flow", 0)
    doc.add_paragraph("Flow Breakdown", style="Heading 1")

    steps = content.split("Step ")
    for step in steps:
        if not step.strip():
            continue

        lines = step.strip().splitlines()
        header_line = lines[0]
        body_lines = lines[1:]

        if ":" in header_line:
            step_num, step_title = header_line.split(":", 1)
        else:
            step_num = header_line
            step_title = ""

        doc.add_heading(
            f"Step {step_num.strip()}: {step_title.strip()}", level=2)

        for line in body_lines:
            line = line.strip()
            if not line:
                continue
            if ":" in line and not line.startswith("<"):
                key, val = line.split(":", 1)
                para = doc.add_paragraph()
                run = para.add_run(f"{key.strip()}: ")
                run.bold = True
                para.add_run(val.strip())
            else:
                doc.add_paragraph(line)

    return doc


def generate_integration_flow_doc(content: str) -> str:
    doc = format_integration_flow_doc("Integration Flow", content)
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    filename = f"integration_flow_{timestamp}.docx"
    output_dir = "generated_docs"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)
    return file_path


@cl.on_chat_start
async def setup():
    global AGENT_LOOKUP
    agents = get_agents_by_name(project_client, AGENT_NAMES)
    AGENT_LOOKUP = {a.name: a for a in agents}

    if not AGENT_LOOKUP:
        await cl.Message("No agents found. Please check your Foundry setup.").send()
        return

    thread = project_client.agents.create_thread()
    cl.user_session.set("thread", thread)

    await cl.Message(
        "Multi-agent session started. Type a prompt to get started."
    ).send()


@cl.on_message
async def on_message(message: cl.Message):
    def parse_handoff(content: str):
        match = re.search(r"@handoff:(.*?):(.*)", content)
        if match:
            return match.group(1).strip(), match.group(2).strip()
        return None, None

    async def agent_to_agent_ask(
        asking_agent_name: str,
        responding_agent_name: str,
        question: str,
        thread,
        wait_time=10
    ):
        asking_agent = AGENT_LOOKUP[asking_agent_name]
        responding_agent = AGENT_LOOKUP[responding_agent_name]

        project_client.agents.create_message(
            thread_id=thread.id, role="user", content=question
        )

        run = project_client.agents.create_and_process_run(
            thread_id=thread.id, agent_id=responding_agent.id
        )
        time.sleep(wait_time)

        messages = project_client.agents.list_messages(thread_id=thread.id)
        sorted_messages = sorted(messages.data, key=lambda x: x.created_at)
        assistant_msg = next(
            (m for m in reversed(sorted_messages) if m.role == "assistant"), None
        )

        if assistant_msg:
            return next(
                (item["text"]["value"]
                 for item in assistant_msg.content if item["type"] == "text"), ""
            )
        return None

    user_input = message.content
    thread = cl.user_session.get("thread")

    if not thread:
        await cl.Message("No active thread.").send()
        return

    agent_name = detect_target_agent(user_input)
    agent = AGENT_LOOKUP.get(agent_name) if agent_name else None

    if not agent:
        await cl.Message(
            "Could not determine which agent to use. Try including keywords like "
            "`requirements`, `architecture`, or `project`."
        ).send()
        return

    try:
        project_client.agents.create_message(
            thread_id=thread.id, role="user", content=user_input
        )
    except Exception as e:
        await cl.Message(f"Failed to send message to thread: {e}").send()
        return

    await cl.Message(f"`{agent.name}` is processing...").send()

    try:
        run = project_client.agents.create_and_process_run(
            thread_id=thread.id, agent_id=agent.id
        )
        time.sleep(10)

        messages = project_client.agents.list_messages(thread_id=thread.id)
        sorted_messages = sorted(messages.data, key=lambda x: x.created_at)
        assistant_msg = next(
            (m for m in reversed(sorted_messages) if m.role == "assistant"), None
        )

        if assistant_msg:
            for item in assistant_msg.content:
                if item["type"] == "text":
                    content = item["text"]["value"]

                    # Check for agent-to-agent handoff
                    target_agent_name, handoff_question = parse_handoff(
                        content)
                    if target_agent_name and handoff_question:
                        await cl.Message(
                            f"🤖 `{agent.name}` is unsure and is asking `{target_agent_name}`:\n\n❓ *{handoff_question}*"
                        ).send()

                        handoff_response = await agent_to_agent_ask(
                            asking_agent_name=agent.name,
                            responding_agent_name=target_agent_name,
                            question=handoff_question,
                            thread=thread
                        )

                        if handoff_response:
                            await cl.Message(
                                f" Response from `{target_agent_name}`:\n\n{handoff_response}"
                            ).send()
                        else:
                            await cl.Message(
                                f"`{target_agent_name}` did not respond to the follow-up."
                            ).send()
                        return

                    # Handle document generation
                    if agent.name == "AIS Developer" and should_generate_flow_doc(user_input):
                        file_path = generate_integration_flow_doc(content)
                        elements = [
                            cl.File(
                                name=os.path.basename(file_path),
                                path=file_path,
                                display_name="Integration Flow"
                            ),
                        ]
                        await cl.Message(
                            author=agent.name,
                            content="Integration flow document generated (click to download):",
                            elements=elements
                        ).send()
                    else:
                        await cl.Message(author=agent.name, content=content).send()
                    return

        await cl.Message("No reply found from the agent.").send()

    except Exception as e:
        await cl.Message(f"Error running `{agent.name}`: {e}").send()
